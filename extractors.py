"""문서(PDF/DOCX) 텍스트 추출기 모음.

각 추출기는 `extract(path) -> ExtractResult` 형태로 동작하며,
PDF와 DOCX를 각각 다른 백엔드로 처리한다.
"""

from __future__ import annotations

import html
import os
import re
import time
import traceback
from dataclasses import dataclass, field
from functools import cached_property
from pathlib import Path

# 화면에 미리보기로 띄울 표의 최대 개수. 이 수만큼만 수집한다.
MAX_TABLES = 2

# unstructured hi_res 가 페이지를 렌더링할 해상도. 기본값 200 에서는 한글 표
# 셀이 "여성아。: 죄 AME" 처럼 깨졌고, 300 에서 글자가 제대로 잡혔다(실측).
HI_RES_DPI = 300

# 표 한 개 = 행들의 리스트, 행 한 개 = 셀 문자열의 리스트
Table = list


# ------------------------------------------------------- 본문 분량 정규화
#
# 방식마다 같은 내용을 담아도 원문 길이가 크게 달라진다. 그대로 세면 분량 비교가
# 무의미해진다 (실측: 같은 PDF 에서 Docling 16,218자 vs PyMuPDF 5,990자인데
# 실제 본문은 양쪽 다 ~4,500자).
#   - Docling 의 Markdown 표는 열 너비를 공백으로 맞춰 채운다. 위 문서에선
#     이 정렬 공백과 `|` 가 전체의 70% 이상을 차지했다.
#   - unstructured 는 표를 <table> HTML 통째로 넣는다 (태그가 곧 글자 수).
#   - 우리가 붙인 `--- [page n] ---` 같은 표시도 방식마다 개수가 다르다.
# 그래서 마크업/스캐폴딩을 걷어낸 "본문 문자"를 따로 세서 이걸 기준으로 비교한다.

_RE_HTML_COMMENT = re.compile(r"<!--.*?-->", re.S)
_RE_HTML_TAG = re.compile(r"</?[A-Za-z][^>]*>")
_RE_SCAFFOLD = re.compile(r"^(?:--- \[page \d+\] ---|\[표 \d+\])\s*$", re.M)
# Markdown 표의 구분선 행: |---|:--:|
_RE_MD_SEP_ROW = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|[\s:|-]*$", re.M)
_RE_MD_PREFIX = re.compile(r"^ {0,3}(?:#{1,6}\s+|[-*+]\s+|\d+\.\s+|>\s?)", re.M)
_RE_MD_IMAGE = re.compile(r"!\[[^\]]*\]\([^)]*\)")
_RE_MD_LINK = re.compile(r"\[([^\]]*)\]\([^)]*\)")
# 강조 기호. `_` 는 snake_case 를 깨뜨리지 않도록 단어 경계에서만 지운다.
_RE_EMPHASIS = re.compile(r"\*\*|__|[*`~]|(?<![\w가-힣])_|_(?![\w가-힣])")


def clean_text(text: str) -> str:
    """마크업·스캐폴딩을 걷어낸 본문. 방식 간 분량 비교의 기준이 된다."""
    t = _RE_HTML_COMMENT.sub(" ", text)
    t = _RE_HTML_TAG.sub(" ", t)
    t = html.unescape(t)
    t = _RE_SCAFFOLD.sub("", t)
    t = _RE_MD_SEP_ROW.sub("", t)
    t = _RE_MD_IMAGE.sub("", t)
    t = _RE_MD_LINK.sub(r"\1", t)
    t = _RE_MD_PREFIX.sub("", t)
    t = t.replace("|", " ")  # 표 셀 구분자
    t = _RE_EMPHASIS.sub("", t)
    # 줄 구조는 남기되(줄 수 비교용) 줄 안의 공백은 하나로 접고 빈 줄은 버린다.
    lines = [" ".join(ln.split()) for ln in t.splitlines()]
    return "\n".join(ln for ln in lines if ln)


@dataclass
class ExtractResult:
    method: str
    backend: str
    text: str = ""
    elapsed: float = 0.0
    is_markdown: bool = False
    meta: dict = field(default_factory=dict)
    error: str | None = None
    tables: list = field(default_factory=list)
    # 표가 안 잡혔을 때 그 이유. 화면에서 "표가 없다"는 말 대신 이걸 띄운다.
    tables_note: str | None = None

    @property
    def ok(self) -> bool:
        return self.error is None

    @cached_property
    def clean(self) -> str:
        return clean_text(self.text)

    @property
    def n_raw_chars(self) -> int:
        """추출 결과 원문 길이 (마크업 포함). 다운로드되는 파일의 크기."""
        return len(self.text)

    @property
    def n_chars(self) -> int:
        """본문 문자 수 (마크업·공백 제외). 방식 간 비교는 이 값으로 한다."""
        return sum(len(ln.replace(" ", "")) for ln in self.clean.splitlines())

    @property
    def n_words(self) -> int:
        return len(self.clean.split())

    @property
    def n_lines(self) -> int:
        return len(self.clean.splitlines())


def _ext(path: str | Path) -> str:
    return Path(path).suffix.lower()


# ---------------------------------------------------------------- 1. Docling

_DOCLING_CONVERTER = None


def _docling_converter():
    global _DOCLING_CONVERTER
    if _DOCLING_CONVERTER is None:
        from docling.document_converter import DocumentConverter

        _DOCLING_CONVERTER = DocumentConverter()
    return _DOCLING_CONVERTER


def extract_docling(path: str | Path) -> ExtractResult:
    """Docling: 레이아웃/표 구조를 인식해 Markdown으로 변환."""
    res = ExtractResult(method="Docling", backend="docling.DocumentConverter", is_markdown=True)
    t0 = time.perf_counter()
    try:
        conv = _docling_converter().convert(str(path))
        doc = conv.document
        res.text = doc.export_to_markdown()
        res.meta = {
            "페이지 수": getattr(doc, "num_pages", lambda: "-")()
            if callable(getattr(doc, "num_pages", None))
            else "-",
            "표 개수": len(getattr(doc, "tables", []) or []),
            "그림 개수": len(getattr(doc, "pictures", []) or []),
        }
        for t in (getattr(doc, "tables", None) or [])[:MAX_TABLES]:
            try:
                df = t.export_to_dataframe(doc)
            except TypeError:  # 구버전 시그니처
                df = t.export_to_dataframe()
            clean = lambda v: " ".join(str(v).split())  # noqa: E731 — 셀 내 줄바꿈 정리
            res.tables.append(
                [[clean(c) for c in df.columns]]
                + [[clean(c) for c in row] for row in df.values.tolist()]
            )
        if not res.tables:
            res.tables_note = (
                "레이아웃 모델이 이 문서에서 표 영역을 찾지 못했습니다. "
                "괘선 없이 공백으로만 정렬된 표는 본문 문단으로 분류될 수 있습니다."
            )
    except Exception as e:  # noqa: BLE001
        res.error = f"{type(e).__name__}: {e}\n\n{traceback.format_exc(limit=3)}"
    res.elapsed = time.perf_counter() - t0
    return res


# ------------------------------------------- 2. PyMuPDF / python-docx (네이티브)


def _extract_pymupdf(path: str | Path, res: ExtractResult) -> None:
    import pymupdf

    parts = []
    with pymupdf.open(str(path)) as doc:
        for i, page in enumerate(doc, 1):
            parts.append(f"--- [page {i}] ---\n{page.get_text('text')}")
            if len(res.tables) < MAX_TABLES:
                for tab in page.find_tables().tables:
                    res.tables.append(
                        [[(c or "").strip() for c in row] for row in tab.extract()]
                    )
                    if len(res.tables) >= MAX_TABLES:
                        break
        res.meta = {"페이지 수": doc.page_count, "메타데이터 제목": doc.metadata.get("title") or "-"}
    res.text = "\n".join(parts)


def _extract_python_docx(path: str | Path, res: ExtractResult) -> None:
    import docx

    d = docx.Document(str(path))
    parts = []
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "1"
            parts.append(f"{'#' * min(int(level), 6)} {p.text}")
        else:
            parts.append(p.text)
    for ti, table in enumerate(d.tables, 1):
        parts.append(f"\n[표 {ti}]")
        for row in table.rows:
            parts.append(" | ".join(c.text.replace("\n", " ").strip() for c in row.cells))
        if len(res.tables) < MAX_TABLES:
            res.tables.append(
                [[c.text.replace("\n", " ").strip() for c in r.cells] for r in table.rows]
            )
    res.text = "\n\n".join(parts)
    res.meta = {"문단 수": len(d.paragraphs), "표 개수": len(d.tables), "섹션 수": len(d.sections)}


def extract_native(path: str | Path) -> ExtractResult:
    """PDF는 PyMuPDF, DOCX는 python-docx로 원문 텍스트를 그대로 추출."""
    ext = _ext(path)
    backend = "PyMuPDF (fitz)" if ext == ".pdf" else "python-docx"
    res = ExtractResult(method="PyMuPDF / python-docx", backend=backend)
    t0 = time.perf_counter()
    try:
        if ext == ".pdf":
            _extract_pymupdf(path, res)
        elif ext == ".docx":
            _extract_python_docx(path, res)
            res.is_markdown = True
        else:
            raise ValueError(f"지원하지 않는 확장자: {ext}")
    except Exception as e:  # noqa: BLE001
        res.error = f"{type(e).__name__}: {e}\n\n{traceback.format_exc(limit=3)}"
    res.elapsed = time.perf_counter() - t0
    return res


# ------------------------------------------------ 3. pdfplumber / mammoth


def _extract_pdfplumber(path: str | Path, res: ExtractResult) -> None:
    import pdfplumber

    parts, n_tables = [], 0
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            parts.append(f"--- [page {i}] ---")
            # extract_text 에는 표 안의 글자도 이미 들어 있다. 여기에 표를 또
            # 덧붙이면 본문이 그대로 두 번 들어가(실측 4,499자 → 8,479자) 분량
            # 비교가 깨지고 RAG 청킹에도 중복이 섞인다. 표는 아래 res.tables 로만
            # 따로 넘겨 미리보기에서 구조를 보여준다.
            parts.append(page.extract_text(layout=False) or "")
            for table in page.extract_tables() or []:
                n_tables += 1
                if len(res.tables) < MAX_TABLES:
                    res.tables.append(
                        [[(c or "").replace("\n", " ").strip() for c in row] for row in table]
                    )
        res.meta = {"페이지 수": len(pdf.pages), "검출된 표": n_tables}
    res.text = "\n".join(parts)
    if not n_tables:
        res.tables_note = (
            "pdfplumber 의 표 검출은 셀 경계선(괘선)에 의존합니다. "
            "괘선 없이 공백으로만 정렬된 표는 검출되지 않고 본문 텍스트로만 남습니다."
        )


def _extract_mammoth(path: str | Path, res: ExtractResult) -> None:
    import mammoth

    with open(path, "rb") as f:
        out = mammoth.convert_to_markdown(f)
    res.text = out.value
    res.is_markdown = True
    res.meta = {"변환 경고": len(out.messages)}
    if out.messages:
        res.meta["경고 예시"] = str(out.messages[0])[:120]


def extract_layout(path: str | Path) -> ExtractResult:
    """PDF는 pdfplumber(표 검출), DOCX는 mammoth(스타일→Markdown)."""
    ext = _ext(path)
    backend = "pdfplumber" if ext == ".pdf" else "mammoth"
    res = ExtractResult(method="pdfplumber / mammoth", backend=backend)
    t0 = time.perf_counter()
    try:
        if ext == ".pdf":
            _extract_pdfplumber(path, res)
        elif ext == ".docx":
            _extract_mammoth(path, res)
        else:
            raise ValueError(f"지원하지 않는 확장자: {ext}")
    except Exception as e:  # noqa: BLE001
        res.error = f"{type(e).__name__}: {e}\n\n{traceback.format_exc(limit=3)}"
    res.elapsed = time.perf_counter() - t0
    return res


# ------------------------------------------------------- 4. unstructured

# unstructured 요소 카테고리 → Markdown 표현
_UNS_MARKUP = {
    "Title": "## {}",
    "Header": "### {}",
    "ListItem": "- {}",
    "Footer": "_{}_",
    "PageBreak": "\n---\n",
}


def _neutralize_broken_libmagic() -> None:
    """로드 불가능한 libmagic 때문에 프로세스가 죽는 것을 막는다.

    Windows에서 python-magic 은 후보를 순서대로 훑다가 마지막에 %PATH% 의
    Git for Windows `msys-magic-1.dll` 을 집는다. 이건 MSYS2 빌드라 네이티브
    파이썬에 ctypes 로 로드되는 순간 access violation 으로 인터프리터가 죽는다
    (try/except 로 잡히지 않는 세그폴트).

    unstructured 는 `magic` 임포트가 실패하면 `filetype` 패키지 기반 탐지로
    폴백하도록 이미 만들어져 있다. 그래서 sys.modules 에 None 을 심어
    `import magic` 을 ImportError 로 만들어 그 폴백 경로로 보낸다.
    정상 libmagic 이 설치된 환경이라면 손대지 않는다.
    """
    import sys

    if sys.platform not in ("win32", "cygwin") or "magic" in sys.modules:
        return

    from ctypes.util import find_library

    # python-magic 이 msys-magic-1 보다 먼저 시도하는 후보들
    for name in ("magic", "libmagic", "magic1", "cygmagic-1", "libmagic-1"):
        if find_library(name) or Path(f"./{name}.dll").exists():
            return  # 쓸 수 있는 libmagic 이 있으니 그대로 둔다

    sys.modules["magic"] = None  # type: ignore[assignment]


def _rows_from_html(html: str | None) -> list:
    """unstructured 가 표를 담아 주는 HTML(<table>)을 행/셀 리스트로 바꾼다."""
    if not html:
        return []
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return []
    soup = BeautifulSoup(html, "html.parser")
    rows = []
    for tr in soup.find_all("tr"):
        cells = [c.get_text(" ", strip=True) for c in tr.find_all(["td", "th"])]
        if cells:
            rows.append(cells)
    return rows


# hi_res 전략이 내부적으로 호출하는 외부 바이너리 (pdftoppm = poppler-utils)
_HI_RES_BINARIES = {"tesseract": "tesseract-ocr", "pdftoppm": "poppler-utils"}

# 문서가 한국어라 kor 을 먼저 준다. 실제로 설치된 것만 골라 쓴다.
_PREFERRED_OCR_LANGS = ("kor", "eng")

# Windows 표준 설치 경로. 설치 관리자가 PATH 를 안 건드리거나(tesseract),
# 건드려도 이미 떠 있는 셸에는 반영되지 않는(poppler) 경우를 메우기 위한 후보들.
_WINDOWS_BIN_HINTS = {
    "tesseract": (
        r"C:\Program Files\Tesseract-OCR",
        r"C:\Program Files (x86)\Tesseract-OCR",
    ),
    "pdftoppm": (
        # winget(oschwartz10612.Poppler) 은 버전이 폴더명에 박히므로 glob 로 찾는다.
        r"%LOCALAPPDATA%\Microsoft\WinGet\Packages\*Poppler*\poppler-*\Library\bin",
        r"%LOCALAPPDATA%\Microsoft\WinGet\Links",
    ),
}


def _prepare_ocr_env() -> None:
    """OCR 바이너리·언어 데이터를 현재 프로세스에서 찾을 수 있게 맞춘다.

    Windows 전용 보정이다. 설치 직후의 PATH/TESSDATA_PREFIX 변경은 이미 떠 있는
    셸에 전달되지 않아, 터미널을 새로 열기 전까지 hi_res 가 조용히 fast 로
    떨어진다. 표준 설치 경로를 직접 확인해 그 구멍을 메운다.
    Linux 배포(apt) 환경은 이미 PATH 에 있으므로 아무 일도 하지 않는다.
    """
    import glob
    import shutil
    import sys

    if sys.platform != "win32":
        return

    for exe, hints in _WINDOWS_BIN_HINTS.items():
        if shutil.which(exe):
            continue
        for hint in hints:
            found = [d for d in glob.glob(os.path.expandvars(hint))
                     if Path(d, f"{exe}.exe").exists()]
            if found:
                os.environ["PATH"] = os.environ.get("PATH", "") + os.pathsep + found[0]
                break

    # 관리자 권한이 없으면 언어 파일을 Program Files 밑에 못 넣는다. 사용자
    # 폴더에 받아 둔 tessdata 가 있으면 그쪽을 가리킨다.
    if not os.environ.get("TESSDATA_PREFIX"):
        user_tessdata = Path(os.environ.get("LOCALAPPDATA", ""), "tessdata")
        if any(user_tessdata.glob("*.traineddata")):
            os.environ["TESSDATA_PREFIX"] = str(user_tessdata)


def _ocr_languages() -> list:
    """설치돼 있는 OCR 언어만 골라 준다 (없는 언어를 넘기면 tesseract 가 죽는다)."""
    try:
        import unstructured_pytesseract

        available = set(unstructured_pytesseract.get_languages())
    except Exception:  # noqa: BLE001 — 못 물어보면 기본값에 맡긴다
        return []
    return [lang for lang in _PREFERRED_OCR_LANGS if lang in available]


def _unstructured_pdf_kwargs(res: ExtractResult) -> dict:
    """PDF 표 인식이 가능한 전략을 고른다.

    unstructured 에서 Table 요소는 레이아웃 모델을 돌리는 `hi_res` 전략에서만
    나온다. `fast`(pdfminer) 는 텍스트 레이어만 읽어서 표를 그냥 문단으로
    취급하므로 `infer_table_structure=True` 를 줘도 Table 요소가 0개다
    (실측 확인). hi_res 는 tesseract·poppler 바이너리를 요구하니, 있으면 쓰고
    없으면 fast 로 내리되 그 이유를 결과에 남긴다.
    """
    import shutil

    _prepare_ocr_env()
    missing = [pkg for exe, pkg in _HI_RES_BINARIES.items() if not shutil.which(exe)]
    if not missing:
        langs = _ocr_languages()
        res.backend = "unstructured (strategy=hi_res)"
        kwargs = {
            "strategy": "hi_res",
            "infer_table_structure": True,
            # 표 셀 내용은 페이지를 렌더링해 OCR 로 읽는다. 기본 200dpi 로는
            # 한글 획이 뭉개져 셀 텍스트가 깨진다.
            "pdf_image_dpi": HI_RES_DPI,
        }
        if langs:
            kwargs["languages"] = langs
            res.backend += f" · OCR {'+'.join(langs)}"
        return kwargs

    res.backend = "unstructured (strategy=fast)"
    res.tables_note = (
        "unstructured 는 PDF 에서 `strategy=\"hi_res\"` 일 때만 Table 요소를 만듭니다. "
        f"hi_res 에 필요한 **{', '.join(missing)}** 가 이 환경에 없어, 텍스트 레이어만 "
        "읽는 `fast`(pdfminer) 전략으로 실행했습니다. fast 는 표를 일반 문단으로 "
        "취급하므로 표 구조가 남지 않습니다 — 내용 자체는 본문에 들어 있습니다."
    )
    return {"strategy": "fast"}


def extract_unstructured(path: str | Path) -> ExtractResult:
    """unstructured: 문서를 의미 단위 요소(Title/NarrativeText/Table…)로 분해."""
    res = ExtractResult(
        method="unstructured", backend="unstructured.partition", is_markdown=True
    )
    t0 = time.perf_counter()
    try:
        _neutralize_broken_libmagic()
        from unstructured.partition.auto import partition

        kwargs = _unstructured_pdf_kwargs(res) if _ext(path) == ".pdf" else {}
        used_hi_res = kwargs.get("strategy") == "hi_res"
        try:
            elements = partition(filename=str(path), **kwargs)
        except Exception:  # noqa: BLE001 — hi_res 는 런타임에도 깨질 수 있다
            if not used_hi_res:
                raise
            used_hi_res = False
            res.backend = "unstructured (strategy=fast, hi_res 실패 후 폴백)"
            res.tables_note = (
                "hi_res 전략이 실행 중 실패해 `fast` 로 다시 추출했습니다. "
                "fast 는 표 구조를 인식하지 않습니다."
            )
            elements = partition(filename=str(path), strategy="fast")

        parts, counts = [], {}
        for el in elements:
            cat = el.category
            counts[cat] = counts.get(cat, 0) + 1
            text = (el.text or "").strip()
            if cat == "Table":
                table_html = getattr(el.metadata, "text_as_html", None)
                # 본문에는 el.text 를 쓴다. text_as_html 의 셀 내용은 표 이미지를
                # 셀 단위로 OCR 한 결과라 한글이 뭉개지는데, el.text 는 PDF 텍스트
                # 레이어에서 온 깨끗한 문장이다 (실측: html 을 쓰면 본문이
                # 4,514자 → 3,201자로 줄고 내용도 망가졌다).
                # 표의 행/열 구조는 아래 res.tables 로만 넘긴다.
                if text:
                    parts.append(text)
                if len(res.tables) < MAX_TABLES:
                    rows = _rows_from_html(table_html) or [
                        [ln] for ln in text.splitlines() if ln
                    ]
                    if rows:
                        res.tables.append(rows)
                continue
            if not text:
                continue
            parts.append(_UNS_MARKUP.get(cat, "{}").format(text))

        res.text = "\n\n".join(parts)
        # 그림 개수는 상위 5개 카테고리에 밀려 빠질 수 있어 따로 못 박아 둔다.
        res.meta = {"요소 수": len(elements), "그림 개수": counts.get("Image", 0)}
        res.meta.update(
            {k: v for k, v in sorted(counts.items(), key=lambda kv: -kv[1])[:5]}
        )
        if not res.tables and not res.tables_note:
            res.tables_note = "이 문서에서는 표로 판정된 영역이 없었습니다."
    except Exception as e:  # noqa: BLE001
        res.error = f"{type(e).__name__}: {e}\n\n{traceback.format_exc(limit=3)}"
    res.elapsed = time.perf_counter() - t0
    return res


EXTRACTORS = {
    "Docling": {
        "fn": extract_docling,
        "desc": "AI 레이아웃 분석 기반. 읽기 순서·표·수식을 인식해 Markdown으로 구조화.",
        "backends": "PDF/DOCX 모두 `docling`",
        # cost: 실행 순서 결정용 예상 비용. 낮을수록 먼저 실행해 먼저 화면에 그린다.
        "cost": 2,
        # images: 그림 영역을 인식하는가. images_note: 그 근거/한계 한 줄.
        "images": "O",
        "images_note": "레이아웃 모델이 그림 영역을 검출해 개수를 세고, 본문에 `<!-- image -->` 로 위치를 남깁니다.",
        # cells: 표의 셀 내용을 원문 그대로 읽는가. 구조는 모델이 예측하더라도
        # 칸을 채우는 글자는 PDF 텍스트 레이어에서 가져오면 O, OCR 로 다시
        # 읽으면 X (한글이 어긋난다).
        "cells": "O",
    },
    "PyMuPDF / python-docx": {
        "fn": extract_native,
        "desc": "포맷 네이티브 파서. 가장 빠르고 원문에 충실하지만 구조 복원은 최소.",
        "backends": "PDF → `pymupdf`, DOCX → `python-docx`",
        "cost": 0,
        # 화면에는 그 파일에 실제로 쓰인 백엔드 하나만 이름으로 띄운다.
        "labels": {".pdf": "PyMuPDF", ".docx": "python-docx"},
        "images": "X",
        "images_note": "텍스트 레이어만 읽습니다. 문서에 박힌 그림은 결과에 흔적조차 남지 않습니다.",
        "cells": "O",
    },
    "pdfplumber / mammoth": {
        "fn": extract_layout,
        "desc": "좌표 기반 표 검출(pdfplumber) / 스타일 기반 Markdown 변환(mammoth).",
        "backends": "PDF → `pdfplumber`, DOCX → `mammoth`",
        "cost": 1,
        "labels": {".pdf": "pdfplumber", ".docx": "mammoth"},
        "images": "X",
        "images_note": "글자와 괘선의 좌표만 봅니다. 그림 영역은 인식 대상이 아닙니다.",
        "cells": "O",
    },
    "unstructured": {
        "fn": extract_unstructured,
        "desc": "문서를 Title/NarrativeText/Table 등 의미 단위 요소로 분해. RAG 청킹에 유리.",
        "backends": "PDF/DOCX 모두 `unstructured` "
        "(PDF는 tesseract·poppler 가 있으면 strategy=hi_res, 없으면 fast)",
        # hi_res 는 페이지를 렌더링해 레이아웃 모델 + OCR 을 돌려 Docling 보다도 느리다.
        "cost": 3,
        "images": "X",
        "images_note": "fast 전략은 텍스트 레이어만 읽습니다. hi_res 로 올라가야 Image 요소가 생깁니다.",
        "images_note_on": "hi_res 레이아웃 모델이 그림 영역을 Image 요소로 분리했습니다.",
        # 셀 내용을 페이지 이미지 OCR 로 읽는 유일한 방식이라 한글이 어긋난다.
        "cells": "X",
    },
}

# 추출기가 meta 에 남기는 "그림 개수" 계열 키 (unstructured 는 카테고리 이름이 Image).
_IMAGE_META_KEYS = ("그림 개수", "Image")


def image_support(method: str, res: ExtractResult | None = None) -> tuple[str, str]:
    """이미지 인식 여부(O/X)와 그 근거 한 줄.

    결과가 있으면 실제로 검출된 그림 개수를 근거에 덧붙인다. 전략에 따라
    달라지는 방식(unstructured hi_res)도 실측값이 있으면 그쪽을 따른다.
    """
    info = EXTRACTORS[method]
    found = 0
    if res is not None and res.ok:
        # 같은 뜻의 키가 둘 다 있을 수 있으니 합치지 말고 큰 쪽을 쓴다.
        found = max(
            (v for k in _IMAGE_META_KEYS if isinstance(v := res.meta.get(k), int)),
            default=0,
        )
    if found:
        # 전략에 따라 되기도 안 되기도 하는 방식은 실제로 됐을 때 쓸 설명이 따로 있다.
        return "O", f"그림 {found}개 검출 · {info.get('images_note_on', info['images_note'])}"
    return info["images"], info["images_note"]


def method_label(method: str, path: str | Path) -> str:
    """해당 파일에 실제로 쓰인 백엔드 이름. 없으면 방식 이름 그대로."""
    return (EXTRACTORS[method].get("labels") or {}).get(_ext(path), method)
